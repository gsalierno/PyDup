from Clusterize import clusterize_questions
from docx import Document

QUEST_PATH = "../questions_raw_cleaned_aggregated.txt"
#get question per topic and questions + answer dictionary
topic_list, questions = clusterize_questions(QUEST_PATH)

#Create document
document = Document()
#adding title
document.add_heading('Domande per il corso di progettazione di sistemi operativi', 0)

for topic in topic_list:
    document.add_heading(topic[0], level=1)
    for q in topic[1]:
        paragraph = document.add_paragraph()
        paragraph.add_run(q+"\n")
        #retrieve answers for q
        answ_list = questions[q]
        for a in answ_list:
            document.add_paragraph(a, style='List Bullet')

document.save('domande-pso.docx')