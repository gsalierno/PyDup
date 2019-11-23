from docx import Document
from itertools import groupby
from operator import itemgetter
import re
import itertools


class Parser:

    def __init__(self, document):
        self.document = self.getText(Document(document))
        self.quaestiones = self.extractQA()
        #self.cleanDictionary()

    def getText(self, document):
        doc = document
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)

        return '\n'.join(fullText)

    def writeDictToFile(self):
        with open("dict.txt", "w+") as f:
            f.write(str(self.quaestiones))
            f.close()

    # extract chunks of QA from unstructured messy text imported from word doc
    def extractQA(self):
        l = []
        q = ""
        for word in self.document:
            if word == '\n':
                l.append(q)
                l.append(word)
                q = ""
            else:
                q += word
        # remove blank lines
        while '' in l:
            l.remove('')

        # aggregate questions and answers
        i = 0
        # retrieve indexes of \n representing possible questions + answers
        indexnl = [i for i, x in enumerate(l) if x == '\n']
        # mantain only consecutive indexes possibily representing begin or end of a Q+A
        range_l = []
        for k, g in groupby(enumerate(indexnl), lambda ix: ix[0] - ix[1]):
            li = list(map(itemgetter(1), g))
            if len(li) == 2:
                range_l.append(li)

        # now we have a list representing possibly begin and ending of questions+answers in format \n\n Q+A \n\n [4, 5], [7, 8], [18, 19], [29, 30], [40, 41], [51, 52], [62, 63]
        # now compute indexes representing begin of q+a and end of q+a
        range_qa = []
        for first, second in zip(range_l, range_l[1:]):
            range_qa.append([first[1], second[1]])
        # try to print questions + answers they as consecutive lists
        qa = []
        for r in range_qa:
            qa.append(l[r[0]:r[1]])

        # FILTERING
        #FILTERING ANSWERS
        a_filtered = self.prefilteringA(qa)
        #FILTERING Q+A
        qa_filtered = self.prefilteringQA(a_filtered)
        #FILTER Q
        q_filtered = self.prefilteringQ(qa_filtered)
        #FILTER A Q
        aq_filtered = self.prefilterAQ((q_filtered))
        # aggregate Q A chunks
        qa_aggr = self.aggregateChunks(aq_filtered)

        for qa in qa_aggr:
            print(qa, "ESTIMATOR: ", self.qaEstimator(qa))

        return self.createQAdictionary(qa_aggr)


    #aggregate chunks of Q with chunks of A
    def aggregateChunks(self, list_qa):
        join = []
        for first, second in zip(list_qa,list_qa[1:]):
            if self.qaEstimator(first) == "Q" and self.qaEstimator(second) == "A":
                join.append((first,second,first+second))

        #remove separate chunks from original list
        for elem in join:
            try:
                index_first = list_qa.index(elem[0]) #retrieve original elem
                index_second =  list_qa.index(elem[1])
                del list_qa[index_first] #remove old chunk q
                del list_qa[index_second] #remove old chunk a
                list_qa.insert(index_first, elem[2]) # insert join of Q+A
            except ValueError:
                print("INDEX NOT FOUND: ", elem, ValueError)
        # delete only chunks representing only answers
        list_qa_clean = [x for x in list_qa if not self.qaEstimator(x) == "A"]

        return list_qa_clean



    def prefilteringQA(self, list_qa):
        replace = []
        for chunk in list_qa:
            if (self.qaEstimator(chunk) == "Q+A"):  # if chunk represent a Q+A
                extracted = self.splitmultipleQA(chunk)
                if (len(extracted) > 0):  # if extract multiple questions store and replace them
                    replace.append((chunk, extracted))

        for elem in replace:
            index = list_qa.index(elem[0]) #get original index to remove and replace
            list_qa.remove(elem[0]) #remove original aggregate question
            for qa in elem[1]: #append newest
                #print(qa)
                list_qa.insert(index, qa)  # remove chunk and add new extracted question and answers
                index = index+1
        return list_qa

    def prefilteringA(self, list_qa):
        replace = []
        for chunk in list_qa:
            if (self.qaEstimator(chunk) == "A"):
                    extracted = self.splitmultipleQA(chunk)
                    if (len(extracted) > 0):  # if extract multiple questions store and replace them
                        replace.append((chunk, extracted))

        for elem in replace:
            index = list_qa.index(elem[0]) #get original index to remove and replace
            list_qa.remove(elem[0]) #remove original aggregate question
            for qa in elem[1]: #append newest
                #print(qa)
                list_qa.insert(index, qa)  # remove chunk and add new extracted question and answers
                index = index+1
        return list_qa

    def prefilteringQ(self, list_qa):
        replace = []
        for chunk in list_qa:
            if (self.qaEstimator(chunk) == "Q"):
                    extracted = self.splitmultipleQA(chunk)
                    if (len(extracted) > 0):  # if extract multiple questions store and replace them
                        replace.append((chunk, extracted))

        for elem in replace:
            index = list_qa.index(elem[0]) #get original index to remove and replace
            list_qa.remove(elem[0]) #remove original aggregate question
            for qa in elem[1]: #append newest
                #print(qa)
                list_qa.insert(index, qa)  # remove chunk and add new extracted question and answers
                index = index+1
        return list_qa

    #split the chunk that aggregates a+q
    def prefilterAQ(self, list_qa):
        list_extracted_index = []
        for chunk in list_qa:
            if(self.qaEstimator(chunk) == "Q+A"):
                occ = 0
                for elem in chunk:
                    if re.match('^[0-9]+.', elem):
                        if(occ < 4):
                            occ = occ + 1
                        else:
                            split_index = chunk.index(elem)
                            list_extracted_index.append((list_qa.index(chunk),chunk[0:split_index],chunk[split_index:])) #store the index for replacement and splitted chunks

                #update original list
                for elem in list_extracted_index:
                    del list_qa[elem[0]] #delete aggregate chunks
                    list_qa.insert(elem[0],elem[1]) #store new chunk
                    list_qa.insert(elem[0] + 1, elem[2]) #store new chunk
        return list_qa

    # split multiple Q+A in different chunks
    def splitmultipleQA(self, chunk):
        qa_split = []
        # re.match('^[0-9]+.', chunk[i])
        for i in range(len(chunk)):
            if re.match('^[0-9]+.', chunk[i]):
                qa_split.append(i)  # store the index of the split
        qa_extracted = []
        # split the chunk considering begin,end of each chunk
        for beg, end in zip(qa_split[:-1], qa_split[1:]):
            qa_extracted.append(chunk[beg:end])
        return qa_extracted

    # given a list determine if it is a question or an answer
    def qaEstimator(self, chunk):
        pm = 0
        # determine if it contains lines for open question
        for elem in chunk:
            if "____________________________________________________" in elem:
                return "Q+A"  # it is a question
            elif re.match('^[0-9]+.', elem):
                pm = pm + 1
        if (len(chunk) > 4 and (re.match('^[A-Z0-9]+.', chunk[1]) or re.match('^[A-Z0-9]+.', chunk[0])) and self.checkNumberofAnswersQ(
                chunk)):  # it is an answer weak condition
            return "Q+A"
        elif pm < 3 and pm != 0:
            return "Q"  # it is a question
        elif len(chunk) == 3:  # it is a question
            return "Q"
        elif len(chunk) > 3:
            return "A"
        elif pm == 0:
            print("Undefinied: ", chunk)
            return "Undef"

    # check number of answer in a chunk used by the estimator to interpret text
    def checkNumberofAnswersQ(self, chunk):
        answ = 0
        for elem in chunk:
            if re.match('^[A-Z0-9]+.', elem):
                answ = answ + 1
        if answ > 4:
            return True
        else:
            return False

    # since some Q+A chunks aggregate different (two) questions and answers this method split them into single chunks
    def filterQAchunk(self, chunk):
        match = re.findall('[0-9]+.', str(chunk))
        # if multiple match exists
        if (len(match) > 0 and len(match) == 2):
            i = 0
            match = 0
            for i in range(len(chunk)):
                if (re.match('[0-9]+.', chunk[i])):
                    match = match + 1  # update match
                    if (match == 2):  # if we encounter the second question split the chunks
                        return chunk[0:i], chunk[i:]  # split into two chunks
        return None, None

    # uniform format for answers
    def cleanDictionary(self):
        for k in self.quaestiones.keys():
            if re.search('\s\t', str(k)):  # if answer contains space and tabs
                k_old = k
                k = re.sub('\s\t', '', str(k))
                # update key
                self.quaestiones[k] = self.quaestiones.pop(k_old)

    def createQAdictionary(self, qa_list):
            qadict = {}
            for elem in qa_list:
                #print(elem, "OUTPUT: ", self.qaEstimator(elem),"\n")
                if (self.qaEstimator(elem) == "Q+A"):
                    self.filterQAchunk(elem)
                    quest, answ = self.splitQA(elem)
                    qadict[quest] = answ
            return qadict

    # Give a chunk Q+A split it
    def splitQA(self, QA):
        # preprocessing and cleaning Q+A
        # remove special characters
        qac = list(filter(lambda a: a != '\n', QA))
        Q = qac[0]  # get answer
        qarr = Q.split('. ')  # ignore numeration
        if (len(qarr) > 1):
            qac[0] = qarr[1]  # set clean question
        return qac[0], qac[1:]

    def printDocument(self):
        # print dictionary
        for k in self.quaestiones.keys():
            print("Question: ", k, " Answer: ", self.quaestiones[k])
